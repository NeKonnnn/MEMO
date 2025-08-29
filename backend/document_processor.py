import os
import tempfile
import docx
import PyPDF2
import openpyxl
import pdfplumber
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document

class DocumentProcessor:
    def __init__(self):
        print("Инициализируем DocumentProcessor...")
        # Инициализация векторного хранилища с пустым набором
        self.documents = []
        self.doc_names = []
        self.embeddings = None
        self.vectorstore = None

        print("DocumentProcessor инициализирован")
        self.init_embeddings()
        
    def init_embeddings(self):
        """Инициализация модели для эмбеддингов"""
        print("Инициализируем модель эмбеддингов...")
        try:
            # Загружаем модель для русского языка
            print("Загружаем модель: sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2")
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
            )
            print("Модель эмбеддингов успешно загружена")
        except Exception as e:
            print(f"Ошибка при загрузке модели эмбеддингов: {str(e)}")
            import traceback
            traceback.print_exc()
            self.embeddings = None
    
    def process_document(self, file_path):
        """Обработка документа в зависимости от его типа"""
        file_extension = os.path.splitext(file_path)[1].lower()
        document_text = ""
        
        try:
            print(f"Обрабатываем документ: {file_path} (тип: {file_extension})")
            
            if file_extension == '.docx':
                document_text = self.extract_text_from_docx(file_path)
            elif file_extension == '.pdf':
                document_text = self.extract_text_from_pdf(file_path)
            elif file_extension in ['.xlsx', '.xls']:
                document_text = self.extract_text_from_excel(file_path)
            elif file_extension == '.txt':
                document_text = self.extract_text_from_txt(file_path)
            elif file_extension in ['.jpg', '.jpeg', '.png', '.webp']:
                document_text = self.extract_text_from_image(file_path)
            else:
                return False, f"Неподдерживаемый формат файла: {file_extension}"
            
            print(f"Извлечено текста: {len(document_text)} символов")
            
            # Добавляем документ в коллекцию
            self.add_document_to_collection(document_text, os.path.basename(file_path))
            print(f"Документ добавлен в коллекцию. Всего документов: {len(self.doc_names)}")
            return True, f"Документ {os.path.basename(file_path)} успешно обработан"
            
        except Exception as e:
            print(f"Ошибка при обработке документа: {str(e)}")
            return False, f"Ошибка при обработке документа: {str(e)}"
    
    def extract_text_from_docx(self, file_path):
        """Извлечение текста из DOCX файла"""
        print(f"Извлекаем текст из DOCX файла: {file_path}")
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        result = "\n".join(full_text)
        print(f"Извлечено {len(result)} символов из DOCX")
        return result
    
    def extract_text_from_pdf(self, file_path):
        """Извлечение текста из PDF файла"""
        print(f"Извлекаем текст из PDF файла: {file_path}")
        text = ""
        
        # Используем PDFPlumber для более точного извлечения текста
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            print(f"PDFPlumber успешно извлек {len(text)} символов")
        except Exception as e:
            print(f"Ошибка при извлечении текста с помощью pdfplumber: {str(e)}")
            
            # Резервный метод с PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text() or ""
                print(f"PyPDF2 успешно извлек {len(text)} символов")
            except Exception as e2:
                print(f"Ошибка при извлечении текста с помощью PyPDF2: {str(e2)}")
                raise
        
        return text
    
    def extract_text_from_excel(self, file_path):
        """Извлечение текста из Excel файла"""
        print(f"Извлекаем текст из Excel файла: {file_path}")
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        text_content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_content.append(f"Лист: {sheet_name}")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        row_values.append(str(cell.value))
                if row_values:
                    text_content.append("\t".join(row_values))
        
        result = "\n".join(text_content)
        print(f"Извлечено {len(result)} символов из Excel")
        return result
    
    def extract_text_from_txt(self, file_path):
        """Извлечение текста из TXT файла"""
        print(f"Извлекаем текст из TXT файла: {file_path}")
        try:
            # Пробуем открыть файл как UTF-8
            with open(file_path, 'r', encoding='utf-8') as file:
                result = file.read()
                print(f"UTF-8 успешно извлек {len(result)} символов")
                return result
        except UnicodeDecodeError:
            # Если не удалось открыть как UTF-8, пробуем другие кодировки
            encodings = ['cp1251', 'latin-1', 'koi8-r']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        result = file.read()
                        print(f"{encoding} успешно извлек {len(result)} символов")
                        return result
                except UnicodeDecodeError:
                    continue
            
            # Если все кодировки не подошли, открываем в бинарном режиме
            with open(file_path, 'rb') as file:
                content = file.read()
                result = str(content)
                print(f"Бинарный режим извлек {len(result)} символов")
                return result

    def extract_text_from_image(self, file_path):
        """Извлечение текста из изображения с помощью OCR"""
        print(f"Извлекаем текст из изображения: {file_path}")
        try:
            # Проверяем наличие библиотеки pytesseract
            import pytesseract
            from PIL import Image
            
            # Открываем изображение с помощью Pillow
            img = Image.open(file_path)
            
            # Извлекаем текст с изображения
            text = pytesseract.image_to_string(img, lang='rus+eng')
            
            # Если текст не извлечен, добавляем описание изображения
            if not text.strip():
                result = f"[Изображение: {os.path.basename(file_path)}. OCR не смог извлечь текст.]"
                print(f"OCR не смог извлечь текст, возвращаем описание: {len(result)} символов")
                return result
            
            print(f"OCR успешно извлек {len(text)} символов")
            return text
        except ImportError:
            # Если pytesseract не установлен, возвращаем информацию о файле
            result = f"[Изображение: {os.path.basename(file_path)}. Для распознавания текста требуется установка pytesseract.]"
            print(f"pytesseract не установлен, возвращаем описание: {len(result)} символов")
            return result
        except Exception as e:
            result = f"[Изображение: {os.path.basename(file_path)}. Ошибка при обработке: {str(e)}]"
            print(f"Ошибка при обработке изображения: {len(result)} символов")
            return result
    
    def add_document_to_collection(self, text, doc_name):
        """Добавление документа в коллекцию и обновление векторного хранилища"""
        print(f"Добавляем документ '{doc_name}' в коллекцию...")
        print(f"Длина текста: {len(text)} символов")
        
        # Разбиваем текст на части (уменьшили размер для экономии токенов)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # Уменьшили с 1000 до 500
            chunk_overlap=100,  # Уменьшили с 200 до 100
            length_function=len,
        )
        
        chunks = text_splitter.split_text(text)
        print(f"Создано чанков: {len(chunks)}")
        
        # Создаем документы для langchain
        langchain_docs = []
        for i, chunk in enumerate(chunks):
            langchain_docs.append(
                Document(
                    page_content=chunk,
                    metadata={"source": doc_name, "chunk": i}
                )
            )
        
        # Добавляем в общий список документов
        self.documents.extend(langchain_docs)
        if doc_name not in self.doc_names:
            self.doc_names.append(doc_name)
        
        print(f"Документ добавлен. Всего документов: {len(self.documents)}, имен: {len(self.doc_names)}")
        
        # Обновляем векторное хранилище
        self.update_vectorstore()
    
    def update_vectorstore(self):
        """Обновление или создание векторного хранилища"""
        print(f"Обновляем векторное хранилище...")
        print(f"Документов для индексации: {len(self.documents)}")
        print(f"Модель эмбеддингов: {self.embeddings is not None}")
        print(f"Текущий vectorstore: {self.vectorstore is not None}")
        
        if not self.documents:
            print("Нет документов для индексации")
            return
        
        if not self.embeddings:
            print("Модель эмбеддингов не инициализирована")
            self.init_embeddings()
            if not self.embeddings:
                print("Не удалось инициализировать модель эмбеддингов")
                return
        
        try:
            # Создаем новое векторное хранилище
            print("Создаем векторное хранилище FAISS...")
            self.vectorstore = FAISS.from_documents(self.documents, self.embeddings)
            print(f"Векторное хранилище успешно обновлено, добавлено {len(self.documents)} чанков")
            print(f"Новый vectorstore доступен: {self.vectorstore is not None}")
        except Exception as e:
            print(f"Ошибка при обновлении векторного хранилища: {str(e)}")
            import traceback
            traceback.print_exc()
    
    def query_documents(self, query, k=2):
        """Поиск релевантных документов по запросу"""
        print(f"Ищем релевантные документы для запроса: '{query}'")
        print(f"Векторное хранилище: {self.vectorstore is not None}")
        
        if not self.vectorstore:
            print("Векторное хранилище не инициализировано или пусто")
            return "Векторное хранилище не инициализировано или пусто"
        
        try:
            print(f"Выполняем поиск с k={k}...")
            docs = self.vectorstore.similarity_search(query, k=k)
            print(f"Найдено документов: {len(docs)}")
            
            results = []
            for doc in docs:
                result = {
                    "content": doc.page_content,
                    "source": doc.metadata.get("source", "Неизвестный источник"),
                    "chunk": doc.metadata.get("chunk", 0)
                }
                results.append(result)
                print(f"Документ: {result['source']}, чанк: {result['chunk']}, длина: {len(result['content'])}")
            
            return results
        except Exception as e:
            print(f"Ошибка при поиске по документам: {str(e)}")
            import traceback
            traceback.print_exc()
            return f"Ошибка при поиске по документам: {str(e)}"
    
    def get_document_list(self):
        """Получение списка загруженных документов"""
        print(f"get_document_list вызван. Документы: {self.doc_names}")
        return self.doc_names
    
    def clear_documents(self):
        """Очистка коллекции документов"""
        print("Очищаем коллекцию документов...")
        self.documents = []
        self.doc_names = []
        self.vectorstore = None
        print("Коллекция документов очищена")
        return "Коллекция документов очищена"
    
    def process_query(self, query, agent_function):
        """Обработка запроса с контекстом документов для LLM"""
        print(f"Обрабатываем запрос: {query}")
        print(f"Векторное хранилище: {self.vectorstore is not None}")
        print(f"Количество документов: {len(self.documents)}")
        print(f"Имена документов: {self.doc_names}")
        
        if not self.vectorstore:
            return "Нет загруженных документов. Пожалуйста, загрузите документы перед выполнением запроса."
        
        try:
            # Получаем релевантные документы
            docs = self.query_documents(query)
            print(f"Найдено релевантных фрагментов: {len(docs) if isinstance(docs, list) else 'ошибка'}")
            
            if isinstance(docs, str):  # Если возникла ошибка
                print(f"Ошибка при поиске документов: {docs}")
                return docs
            
            # Формируем контекст из найденных документов
            context = "Контекст из документов:\n\n"
            for i, doc in enumerate(docs):
                context += f"Фрагмент {i+1} (из документа '{doc['source']}'):\n{doc['content']}\n\n"
            
            print(f"Контекст сформирован, длина: {len(context)} символов")
            
            # Подготавливаем запрос для LLM с инструкциями и контекстом
            prompt = f"""На основе предоставленного контекста ответь на вопрос пользователя. 
Если информации в контексте недостаточно, укажи это.
Отвечай только на основе информации из контекста. Не придумывай информацию.

{context}

Вопрос пользователя: {query}

Ответ:"""
            
            print("Отправляем запрос к LLM...")
            # Отправляем запрос к LLM
            response = agent_function(prompt)
            print(f"Получен ответ от LLM, длина: {len(response)} символов")
            return response
            
        except Exception as e:
            print(f"Ошибка при обработке запроса: {str(e)}")
            return f"Ошибка при обработке запроса: {str(e)}"
    
    def remove_document(self, filename):
        """Удалить конкретный документ по имени файла"""
        print(f"Удаляем документ: {filename}")
        print(f"До удаления - self.doc_names: {self.doc_names}")
        print(f"До удаления - self.documents: {len(self.documents)}")
        print(f"До удаления - self.vectorstore доступен: {self.vectorstore is not None}")
        
        try:
            # Находим индекс документа
            if filename not in self.doc_names:
                print(f"Документ {filename} не найден")
                return False
            
            # Удаляем документ из списка имен
            index = self.doc_names.index(filename)
            self.doc_names.pop(index)
            print(f"Документ {filename} удален из списка имен")
            
            # Удаляем ВСЕ чанки этого документа из списка документов
            # Ищем все документы с этим именем и удаляем их
            documents_to_remove = []
            for i, doc in enumerate(self.documents):
                if doc.metadata.get("source") == filename:
                    documents_to_remove.append(i)
            
            # Удаляем чанки в обратном порядке, чтобы индексы не сдвигались
            for i in reversed(documents_to_remove):
                self.documents.pop(i)
            
            print(f"Удалено чанков документа {filename}: {len(documents_to_remove)}")
            print(f"После удаления - self.doc_names: {self.doc_names}")
            print(f"После удаления - self.documents: {len(self.documents)}")
            
            # Пересоздаем vectorstore с оставшимися документами
            if self.documents:
                print("Пересоздаем vectorstore с оставшимися документами")
                self.update_vectorstore()
                print(f"После обновления vectorstore - self.vectorstore доступен: {self.vectorstore is not None}")
            else:
                print("Нет документов, очищаем vectorstore")
                self.vectorstore = None
                print(f"После очистки - self.vectorstore доступен: {self.vectorstore is not None}")
            
            print(f"Документ {filename} успешно удален. Осталось документов: {len(self.doc_names)}")
            return True
            
        except Exception as e:
            print(f"Ошибка при удалении документа {filename}: {str(e)}")
            import traceback
            traceback.print_exc()
            return False 
    
    def get_document_context(self, query, k=2):
        """Получение контекста документов для запроса"""
        print(f"Получаем контекст документов для запроса: '{query}'")
        print(f"Векторное хранилище: {self.vectorstore is not None}")
        
        if not self.vectorstore:
            print("Векторное хранилище не инициализировано")
            return None
        
        try:
            # Получаем релевантные документы
            docs = self.query_documents(query, k=k)
            print(f"Найдено релевантных фрагментов: {len(docs) if isinstance(docs, list) else 'ошибка'}")
            
            if isinstance(docs, str):  # Если возникла ошибка
                print(f"Ошибка при поиске документов: {docs}")
                return None
            
            # Формируем контекст из найденных документов
            context = ""
            for i, doc in enumerate(docs):
                context += f"Фрагмент {i+1} (из документа '{doc['source']}'):\n{doc['content']}\n\n"
            
            print(f"Контекст сформирован, длина: {len(context)} символов")
            
            return context
            
        except Exception as e:
            print(f"Ошибка при получении контекста документов: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
